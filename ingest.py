"""
ingest.py  –  NovaCart RAG Ingestion Pipeline
Reads the 3 source documents, chunks them, embeds and stores in ChromaDB.
Run once before starting the app:  python ingest.py
"""

import os
import re
import json
import hashlib

import docx
import openpyxl
import chromadb
from chromadb.utils import embedding_functions

BASE_DIR    = os.path.dirname(os.path.abspath(__file__))
DATA_DIR    = os.path.join(BASE_DIR, "data")
CHROMA_DIR  = os.path.join(BASE_DIR, "chroma_db")
COLLECTION  = "novacart_knowledge"
EMBED_MODEL = "all-MiniLM-L6-v2"

DOCS = {
    "company_profile": os.path.join(DATA_DIR, "NovaCart_Company_Profile.docx"),
    "qa_knowledge":    os.path.join(DATA_DIR, "NovaCart_Customer_QA_Knowledge_Base.docx"),
    "shipments":       os.path.join(DATA_DIR, "NovaCart_Shipment_Status_Database.xlsx"),
}

CHUNK_SIZE    = 600
CHUNK_OVERLAP = 100


# ── Text helpers ───────────────────────────────────────────────────────────────

def chunk_text(text: str, source: str) -> list:
    text   = re.sub(r"\n{3,}", "\n\n", text.strip())
    chunks = []
    start  = 0
    while start < len(text):
        end   = start + CHUNK_SIZE
        chunk = text[start:end].strip()
        if chunk:
            uid = hashlib.md5(f"{source}::{start}".encode()).hexdigest()
            chunks.append({"id": uid, "text": chunk, "source": source})
        start += CHUNK_SIZE - CHUNK_OVERLAP
    return chunks


def extract_docx(path: str) -> str:
    doc   = docx.Document(path)
    lines = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            lines.append(t)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                lines.append(" | ".join(cells))
    return "\n".join(lines)


def extract_shipments(path: str) -> list:
    wb   = openpyxl.load_workbook(path, data_only=True)
    ws   = wb["Shipments"]
    rows = list(ws.iter_rows(values_only=True))

    header_row = None
    header_idx = 0
    for i, row in enumerate(rows):
        if row and any(str(c).strip() == "Shipment ID" for c in row if c):
            header_row = [str(c).strip() if c else "" for c in row]
            header_idx = i
            break
    if not header_row:
        return []

    records = []
    for row in rows[header_idx + 1:]:
        if not row or not row[0]:
            continue
        rec = dict(zip(header_row, [str(c).strip() if c else "" for c in row]))
        if rec.get("Shipment ID", "").startswith("NVC"):
            records.append(rec)
    return records


def shipment_to_text(rec: dict) -> str:
    return (
        f"Shipment ID: {rec['Shipment ID']} | "
        f"Order ID: {rec['Order ID']} | "
        f"Customer ID: {rec['Customer ID']} | "
        f"Destination: {rec['Destination City']}, {rec['Country']} | "
        f"Courier: {rec['Courier']} | "
        f"Fulfillment Center: {rec['Fulfillment Center']} | "
        f"Created: {rec['Created Date']} | "
        f"Estimated Delivery: {rec['Estimated Delivery']} | "
        f"Status: {rec['Status']} | "
        f"Last Update: {rec['Last Update']} | "
        f"Remarks: {rec['Remarks']}"
    )


# ── Main ───────────────────────────────────────────────────────────────────────

def ingest():
    print("Initializing ChromaDB + embedding model (downloading if needed)...")
    client = chromadb.PersistentClient(path=CHROMA_DIR)
    ef     = embedding_functions.SentenceTransformerEmbeddingFunction(
                 model_name=EMBED_MODEL
             )

    try:
        client.delete_collection(COLLECTION)
        print("Existing collection cleared.")
    except Exception:
        pass

    col = client.create_collection(COLLECTION, embedding_function=ef)

    all_chunks = []

    print("Reading Company Profile...")
    text = extract_docx(DOCS["company_profile"])
    all_chunks.extend(chunk_text(text, "company_profile"))

    print("Reading Q&A Knowledge Base...")
    text = extract_docx(DOCS["qa_knowledge"])
    all_chunks.extend(chunk_text(text, "qa_knowledge"))

    print("Reading Shipment Database...")
    records = extract_shipments(DOCS["shipments"])
    for rec in records:
        t   = shipment_to_text(rec)
        uid = hashlib.md5(rec["Shipment ID"].encode()).hexdigest()
        all_chunks.append({"id": uid, "text": t, "source": "shipments"})

    # Save shipments as JSON for fast exact-match lookup
    ships_path = os.path.join(BASE_DIR, "shipments.json")
    with open(ships_path, "w") as f:
        json.dump({r["Shipment ID"]: r for r in records}, f, indent=2)
    print(f"Saved {len(records)} shipment records to shipments.json")

    print(f"Indexing {len(all_chunks)} chunks into ChromaDB...")
    batch_size = 100
    for i in range(0, len(all_chunks), batch_size):
        batch = all_chunks[i: i + batch_size]
        col.add(
            ids       = [c["id"] for c in batch],
            documents = [c["text"] for c in batch],
            metadatas = [{"source": c["source"]} for c in batch],
        )
        print(f"  {min(i + batch_size, len(all_chunks))}/{len(all_chunks)} chunks indexed")

    print("\nIngestion complete! You can now run:  streamlit run app.py")


if __name__ == "__main__":
    ingest()
